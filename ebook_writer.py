import streamlit as st
import openai
import docx
from io import BytesIO
import requests

# OpenAI API Key (Ensure to set this securely in a real app)
openai.api_key = "your_openai_api_key"


def check_topic_popularity(topic):
    """Checks the popularity of a book topic using a web search API."""
    search_url = f"https://api.example.com/search?query={topic}"  # Replace with actual API
    response = requests.get(search_url)
    if response.status_code == 200:
        popularity_score = response.json().get("popularity_score", 0)
        return popularity_score
    return None


def generate_outline(book_topic, audience):
    """Generates an outline for the book based on the topic and target audience."""
    prompt = f"""
    Generate a detailed chapter outline for an eBook titled '{book_topic}'.
    The book is targeted at {audience} and should be well-structured to maximize reader engagement and marketability.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are an expert book writer and publisher."},
                  {"role": "user", "content": prompt}],
        temperature=0.7
    )
    return response['choices'][0]['message']['content']


def write_chapter(chapter_title, book_topic, audience, previous_content=""):
    """Generates content for a specific chapter."""
    prompt = f"""
    Write a compelling chapter titled '{chapter_title}' for the book '{book_topic}'.
    The book is designed for {audience} and should be engaging, informative, and valuable.
    Ensure it follows a structured format with an introduction, main content, actionable steps, and key takeaways.
    """
    if previous_content:
        prompt += f" Consider the previous content: {previous_content[:500]}..."

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are a professional book author."},
                  {"role": "user", "content": prompt}],
        temperature=0.7
    )
    return response['choices'][0]['message']['content']


def generate_docx(book_title, chapters):
    """Generates a downloadable Word document with the book content."""
    doc = docx.Document()
    doc.add_heading(book_title, 0)

    for chapter_title, content in chapters.items():
        doc.add_heading(chapter_title, level=1)
        doc.add_paragraph(content)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def main():
    st.title("AI eBook Writer")
    st.write("Generate and publish your eBook with AI-powered assistance!")

    book_title = st.text_input("Enter your book idea:")
    audience = st.selectbox("Choose your target audience:",
                            ["Beginners", "Professionals", "Entrepreneurs", "General Readers"])

    if st.button("Check Topic Popularity"):
        popularity = check_topic_popularity(book_title)
        if popularity is not None:
            st.write(f"Popularity Score: {popularity} (Higher is better)")
        else:
            st.write("Could not determine popularity. Proceeding with book outline.")

    if st.button("Generate Outline"):
        outline = generate_outline(book_title, audience)
        st.session_state["outline"] = outline
        st.write("## Suggested Outline:")
        st.write(outline)

    if "outline" in st.session_state:
        st.write("### Proceed with writing chapters")
        chapter_list = st.session_state["outline"].split("\n")
        chapters = {}

        for chapter in chapter_list:
            if chapter.strip():
                if st.button(f"Generate {chapter.strip()}"):
                    chapter_content = write_chapter(chapter.strip(), book_title, audience)
                    chapters[chapter.strip()] = chapter_content
                    st.write(f"### {chapter.strip()}")
                    st.write(chapter_content)

        st.session_state["chapters"] = chapters

        if len(chapters) == len(chapter_list):
            st.write("## Generate Final eBook")
            if st.button("Download eBook as Word Document"):
                doc_buffer = generate_docx(book_title, chapters)
                st.download_button(label="Download eBook", data=doc_buffer, file_name=f"{book_title}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    main()
